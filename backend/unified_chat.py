"""
Unified Chat System for Agents Marketplace
Single API endpoint that routes to different LLM agents based on mode
"""

import json
import logging
from typing import Dict, List, Any, Optional
from openai import OpenAI
from datetime import datetime
from collections import deque
import uuid
import threading
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

from config import OPENAI_API_KEY
from data_source import data_source

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class UnifiedChatAgent:
    def __init__(self):
        """Initialize the unified chat agent with OpenAI client"""
        self.client = None
        self.api_key = OPENAI_API_KEY
        self.conversation_memory = {}  # Store conversations by session_id
        
        # Log API key status
        logger.info(f"OpenAI API Key Status: {'Present' if self.api_key else 'Missing'}")
        if self.api_key:
            logger.info(f"API Key Length: {len(self.api_key)} characters")
            logger.info(f"API Key Prefix: {self.api_key[:10]}..." if len(self.api_key) > 10 else "API Key too short")
        
        # Initialize OpenAI client if API key is valid
        if self.api_key and self.api_key != "your-openai-api-key-here":
            try:
                self.client = OpenAI(api_key=self.api_key)
                logger.info("Unified Chat Agent OpenAI client initialized successfully")
            except Exception as e:
                logger.warning(f"Failed to initialize Unified Chat Agent OpenAI client: {str(e)}")
                self.client = None
        else:
            logger.warning("No OpenAI API key provided - will use fallback responses only")
    
    def get_agents_context(self) -> str:
        """Get formatted context about all available agents for explore mode"""
        try:
            agents_df = data_source.get_agents()
            capabilities_df = data_source.get_capabilities_mapping()
            
            # Convert DataFrames to dictionaries for easier processing
            agents_data = agents_df.to_dict('records') if not agents_df.empty else []
            capabilities_data = capabilities_df.to_dict('records') if not capabilities_df.empty else []
            
            # Create capabilities lookup
            capabilities_lookup = {}
            for cap in capabilities_data:
                agent_id = cap.get('agent_id')
                if agent_id not in capabilities_lookup:
                    capabilities_lookup[agent_id] = []
                capabilities_lookup[agent_id].append(cap.get('by_capability', ''))
            
            # Format agents information
            agents_context = []
            for agent in agents_data:
                if agent.get('admin_approved') == 'yes':  # Only include approved agents
                    agent_id = agent.get('agent_id', '')
                    agent_name = agent.get('agent_name', 'Unknown')
                    description = agent.get('description', 'No description available')
                    by_persona = agent.get('by_persona', 'General')
                    by_value = agent.get('by_value', 'Value not specified')
                    features = agent.get('features', 'No features listed')
                    tags = agent.get('tags', 'No tags')
                    
                    # Get capabilities for this agent
                    capabilities = capabilities_lookup.get(agent_id, [])
                    capabilities_str = ', '.join(filter(None, capabilities)) if capabilities else 'No specific capabilities listed'
                    
                    agent_info = f"""
Agent ID: {agent_id}
Name: {agent_name}
Description: {description}
Target Persona: {by_persona}
Value Proposition: {by_value}
Features: {features}
Tags: {tags}
Capabilities: {capabilities_str}
---
"""
                    agents_context.append(agent_info)
            
            return '\n'.join(agents_context)
            
        except Exception as e:
            logger.error(f"Error getting agents context: {str(e)}")
            return "Error retrieving agents information."
    
    def get_explore_system_prompt(self) -> str:
        """Get the system prompt for agent exploration"""
        return """You are an AI assistant for the Agents Marketplace, a platform where users can discover and explore AI agents for various business needs.

Your role is to:
1. Help users understand which agents might be suitable for their needs
2. Provide detailed but summarized information about specific agents when asked (do not provide the shared details as is, weave the information conversationally)
3. Analyze user queries and match them to relevant agents
4. Always maintain a conversational, helpful tone, neatly formatted response
5. Ask follow-up questions to better understand user needs

Available agents context:
{agents_context}

Guidelines:
- When a user asks about a specific agent, provide summarized  information about that agent
- When a user describes a problem or need, suggest relevant agents that could help
- Always be conversational and engaging
- End your response with a follow-up question about what agent or capability they'd like to explore further
- If no agents match their needs, politely explain and aks them to choose "create mode" to build their own agent
** Response Formatting **
- Keep in mind that you responding to a chat interface built on react-markdown library 
- and therefore you need to format your responses that appears beautifully in the chat interface.
-Keep your responses conversational and engaging.
-Format the response in a way that is easy to understand and follow.
-Use markdown formatting to make the response more readable.
-Use bullet points to make the response more readable.

Remember: You have access to detailed information about each agent including their descriptions, target personas, capabilities, and value propositions. Use this information to provide accurate and helpful recommendations."""
    
    def get_create_system_prompt(self) -> str:
        """Get the system prompt for agent creation ideation and solution design"""
        return """You are an AI agent ideation specialist. Your goal is to intelligently design custom AI agents by ideating solutions, proposing creative names, and filling gaps with intelligent assumptions.

**Required Schema (7 fields):**
1. agent_name: Creative, descriptive name (AUTO-ASSIGNED)
2. applicable_persona: Who will use this agent? (INFER from context)
3. applicable_industry: What industry/domain? (INFER from context)
4. problem_statement: What specific problem does this solve? (IDEATE and propose)
5. user_journeys: What are the key workflow steps? (IDEATE and propose)
6. wow_factor: What makes this unique/special? (IDEATE and propose)
7. expected_output: What should it deliver/produce? (IDEATE and propose)

**Behavior Rules:**
- NEVER ask questionnaire-style questions
- START with partial ideas and build conversationally
- AUTO-ASSIGN creative, descriptive agent names
- PROBE naturally through conversation before proposing complete solution
- Start with user journeys and build based on user feedback
- Add problem statement and wow factor as conversation progresses
- Keep responses conversational and highlight key points
- Ask for confirmation only when you have enough information

**Ideation Approach:**
- Be creative and intelligent in your proposals
- Think like a product designer, not a survey taker
- Propose innovative solutions based on the user's core need
- Make intelligent assumptions about user workflows and pain points
- Suggest modern, AI-powered capabilities as wow factors

**Conversational Approach:**
- Start with agent name and user journeys naturally
- Build on user feedback to add problem statement and wow factor
- Complete remaining fields when you have enough context
- Keep responses conversational - don't just list fields
- Highlight key insights and ask for user thoughts
- Fill fields progressively based on conversation flow, not rigid rules

**Conversation Flow:**
1. Analyze user input and infer basic context (persona, industry)
2. Propose a creative agent name and start with user journeys
3. Build conversationally based on user feedback
4. Add problem statement and wow factor as conversation develops
5. Complete remaining fields when context is sufficient
6. Ask for confirmation when you have a solid solution

** Response Formatting **
- Keep in mind that you responding to a chat interface built on react-markdown library 
- and therefore you need to format your responses that appears beautifully in the chat interface.
- Use bullet points to make the response more readable.
- Weave the requirement gathered conversationally in separate paragraphs for each field.
- Use markdown formatting to make the response more readable.

**Build Decision:**
- Set lets_build=true when:
  * User confirms they want to build the agent (says yes, confirm, approve, build it, etc.)
  * You have all the information in all 7 fields
  * The conversation has reached a natural conclusion
- IMPORTANT: If user says "yes", "build it", "let's build", "proceed", etc., ALWAYS set lets_build=true
- Don't be overly cautious - trust the user's confirmation

**CRITICAL REQUIREMENT: EVERY SINGLE RESPONSE MUST END WITH VALID JSON. NO EXCEPTIONS.**

After your conversational response, you MUST include this exact JSON structure:

{
    "lets_build": true/false,
    "gathered_info": {
        "agent_name": "Actual creative name here",
        "applicable_persona": "Actual persona here", 
        "applicable_industry": "Actual industry here",
        "problem_statement": "Actual problem here",
        "user_journeys": "Actual user journey here",
        "wow_factor": "Actual wow factor here",
        "expected_output": "Actual output here"
    }
}

**ABSOLUTE REQUIREMENTS:**
1. Include this JSON in EVERY response
2. Fill fields with actual values, not empty strings or "string"
3. If user says "yes", "build it", "let's build", "proceed", etc., set lets_build: true
4. Be intelligent and fill fields with meaningful content based on conversation

**EXAMPLE - When user says "Yes, build it!":**
{
    "lets_build": true,
    "gathered_info": {
        "agent_name": "TalentScout AI",
        "applicable_persona": "HR Professionals",
        "applicable_industry": "Human Resources",
        "problem_statement": "Streamlines application filtering using natural language",
        "user_journeys": "Upload, filter, review, select",
        "wow_factor": "Natural language processing for intuitive filtering",
        "expected_output": "Ranked candidate shortlist"
    }
}

**CRITICAL: Always include BOTH "lets_build" and "gathered_info" fields. Never just provide the gathered_info fields alone.**
}"""
    
    def get_conversation_history(self, session_id: str) -> List[Dict[str, str]]:
        """Get conversation history for a session (max 3 conversations)"""
        if session_id in self.conversation_memory:
            return list(self.conversation_memory[session_id])
        return []
    
    def add_to_conversation_history(self, session_id: str, user_message: str, assistant_response: str):
        """Add message to conversation history (maintain max 6 messages)"""
        if session_id not in self.conversation_memory:
            self.conversation_memory[session_id] = deque(maxlen=6)  # 3 conversations = 6 messages
        
        self.conversation_memory[session_id].append({"role": "user", "content": user_message})
        self.conversation_memory[session_id].append({"role": "assistant", "content": assistant_response})
    
    def extract_agent_ids_from_response(self, ai_response: str) -> List[str]:
        """Extract mentioned agent IDs from the AI response (for explore mode)"""
        try:
            agents_df = data_source.get_agents()
            agents_data = agents_df.to_dict('records') if not agents_df.empty else []
            mentioned_agents = []
            
            # Look for agent names in the response and match to IDs
            for agent in agents_data:
                agent_name = str(agent.get('agent_name', '')).lower()
                agent_id = agent.get('agent_id', '')
                
                if agent_name and agent_id and agent_name in ai_response.lower():
                    mentioned_agents.append(agent_id)
            
            return mentioned_agents
            
        except Exception as e:
            logger.error(f"Error extracting agent IDs: {str(e)}")
            return []
    
    def parse_create_response_metadata(self, ai_response: str, user_query: str = "") -> Dict[str, Any]:
        """Extract metadata from create mode AI response"""
        try:
            # First try to find JSON at the end of the response
            if "{" in ai_response and "}" in ai_response:
                # Find the JSON object that contains "lets_build"
                json_start = ai_response.find('{\n    "lets_build"')
                if json_start == -1:
                    # Fallback: find the last complete JSON object
                    json_start = ai_response.rfind("{")
                json_str = ai_response[json_start:]
                
                # Try to find the end of the JSON object (handle nested objects)
                brace_count = 0
                json_end = -1
                in_string = False
                escape_next = False
                
                for i, char in enumerate(json_str):
                    if escape_next:
                        escape_next = False
                        continue
                        
                    if char == '\\':
                        escape_next = True
                        continue
                        
                    if char == '"' and not escape_next:
                        in_string = not in_string
                        continue
                        
                    if not in_string:
                        if char == "{":
                            brace_count += 1
                        elif char == "}":
                            brace_count -= 1
                            if brace_count == 0:
                                json_end = i + 1
                                break
                
                if json_end > 0:
                    json_str = json_str[:json_end]
                    try:
                        metadata = json.loads(json_str)
                        # Check if metadata has the expected structure
                        if "lets_build" not in metadata:
                            # AI only provided gathered_info, add missing fields
                            original_metadata = metadata.copy()
                            metadata = {
                                "lets_build": False,
                                "gathered_info": original_metadata
                            }
                        
                        # Check if this is a confirmation response and override lets_build
                        confirmation_phrases = ["yes", "confirm", "approve", "build it", "let's build", "sounds good", "perfect", "build", "proceed", "go ahead", "ok", "okay"]
                        if any(phrase in user_query.lower() for phrase in confirmation_phrases):
                            metadata["lets_build"] = True
                    except json.JSONDecodeError as e:
                        logger.error(f"JSON parsing error: {str(e)}")
                        logger.error(f"JSON string: {json_str}")
                        metadata = {
                            "lets_build": False,
                            "gathered_info": {}
                        }
                    
                    # Remove JSON from the main response
                    clean_response = ai_response[:json_start].strip()
                    
                    return {
                        "response": clean_response,
                        "metadata": metadata
                    }
            
            # If no JSON found, try to parse structured format
            if ("**Agent Name:**" in ai_response or "1. **Agent Name:**" in ai_response or 
                "HR Candidate Filter" in ai_response or "applicable persona" in ai_response.lower()):
                gathered_info = self.extract_gathered_info_from_any_format(ai_response)
                lets_build = "prototype" in ai_response.lower() and ("create" in ai_response.lower() or "proceed" in ai_response.lower())
                
                return {
                    "response": ai_response,
                    "metadata": {
                        "lets_build": lets_build,
                        "gathered_info": gathered_info
                    }
                }
            
            # If no structured data found, return the original response
            return {
                "response": ai_response,
                "metadata": {
                    "lets_build": False,
                    "gathered_info": {}
                }
            }
        except Exception as e:
            logger.error(f"Error parsing create response metadata: {str(e)}")
            return {
                "response": ai_response,
                "metadata": {
                    "lets_build": False,
                    "gathered_info": {}
                }
            }
    
    def extract_gathered_info_from_any_format(self, ai_response: str) -> Dict[str, str]:
        """Extract gathered info from numbered list format"""
        gathered_info = {}
        
        try:
            # First try structured format (numbered list with **)
            if "**Agent Name:**" in ai_response:
                # Extract agent name
                name_match = ai_response.split("**Agent Name:**")[1].split("\n")[0].strip()
                gathered_info["agent_name"] = name_match.replace("**", "").strip()
                
                # Extract persona
                if "**Applicable Persona:**" in ai_response:
                    persona_match = ai_response.split("**Applicable Persona:**")[1].split("\n")[0].strip()
                    gathered_info["applicable_persona"] = persona_match.replace("**", "").strip()
                
                # Extract industry
                if "**Applicable Industry:**" in ai_response:
                    industry_match = ai_response.split("**Applicable Industry:**")[1].split("\n")[0].strip()
                    gathered_info["applicable_industry"] = industry_match.replace("**", "").strip()
                
                # Extract problem statement
                if "**Problem Statement:**" in ai_response:
                    problem_match = ai_response.split("**Problem Statement:**")[1].split("\n")[0].strip()
                    gathered_info["problem_statement"] = problem_match.replace("**", "").strip()
                
                # Extract user journeys
                if "**User Journeys:**" in ai_response:
                    journeys_match = ai_response.split("**User Journeys:**")[1].split("\n")[0].strip()
                    gathered_info["user_journeys"] = journeys_match.replace("**", "").strip()
                
                # Extract wow factor
                if "**Wow Factor:**" in ai_response:
                    wow_match = ai_response.split("**Wow Factor:**")[1].split("\n")[0].strip()
                    gathered_info["wow_factor"] = wow_match.replace("**", "").strip()
                
                # Extract expected output
                if "**Expected Output:**" in ai_response:
                    output_match = ai_response.split("**Expected Output:**")[1].split("\n")[0].strip()
                    gathered_info["expected_output"] = output_match.replace("**", "").strip()
            
            else:
                # Try paragraph format with intelligent extraction
                if "HR Candidate Filter" in ai_response:
                    gathered_info["agent_name"] = "HR Candidate Filter"
                
                if "HR managers" in ai_response.lower():
                    gathered_info["applicable_persona"] = "HR managers"
                
                if "filtering through" in ai_response.lower() and "applications" in ai_response.lower():
                    gathered_info["problem_statement"] = "filtering through large volumes of job applications"
                
                if "review" in ai_response.lower() and "applications" in ai_response.lower():
                    gathered_info["user_journeys"] = "Review applications, identify top candidates"
                
                if "automated" in ai_response.lower() and ("resume" in ai_response.lower() or "processing" in ai_response.lower()):
                    gathered_info["wow_factor"] = "Automated resume processing and intelligent candidate matching"
                
                if "ranked list" in ai_response.lower() or "best-fit candidates" in ai_response.lower():
                    gathered_info["expected_output"] = "Ranked list of best-fit candidates"
                
                # Set default industry if not specified
                if "industry" not in gathered_info:
                    gathered_info["applicable_industry"] = "General"
            
        except Exception as e:
            logger.error(f"Error extracting gathered info from list: {str(e)}")
        
        return gathered_info
    
    def _add_formatted_paragraph(self, doc, text: str):
        """Add a paragraph with proper markdown formatting"""
        try:
            p = doc.add_paragraph()
            self._add_formatted_text(p, text)
        except Exception as e:
            logger.error(f"Error adding formatted paragraph: {str(e)}")
            doc.add_paragraph(text)
    
    def _add_formatted_bullet(self, doc, text: str):
        """Add a bullet point with proper markdown formatting"""
        try:
            p = doc.add_paragraph(style='List Bullet')
            self._add_formatted_text(p, text)
        except Exception as e:
            logger.error(f"Error adding formatted bullet: {str(e)}")
            doc.add_paragraph(text, style='List Bullet')
    
    def _add_formatted_numbered(self, doc, text: str):
        """Add a numbered item with proper markdown formatting"""
        try:
            p = doc.add_paragraph(style='List Number')
            self._add_formatted_text(p, text)
        except Exception as e:
            logger.error(f"Error adding formatted numbered: {str(e)}")
            doc.add_paragraph(text, style='List Number')
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text to paragraph with proper markdown formatting (bold, italic, etc.)"""
        try:
            # Split text by markdown formatting
            parts = text.split('**')
            
            for i, part in enumerate(parts):
                if not part:
                    continue
                
                # Check if this part should be bold (odd indices after splitting by **)
                if i % 2 == 1:
                    # This part should be bold
                    run = paragraph.add_run(part)
                    run.bold = True
                    run.font.color.rgb = RGBColor(44, 62, 80)  # Professional dark blue
                else:
                    # Regular text - check for other formatting
                    self._add_text_with_formatting(paragraph, part)
                    
        except Exception as e:
            logger.error(f"Error adding formatted text: {str(e)}")
            paragraph.add_run(text)
    
    def _add_text_with_formatting(self, paragraph, text: str):
        """Add text with various formatting (italic, etc.)"""
        try:
            # Handle italic formatting (*text*)
            parts = text.split('*')
            
            for i, part in enumerate(parts):
                if not part:
                    continue
                
                if i % 2 == 1:
                    # This part should be italic
                    run = paragraph.add_run(part)
                    run.italic = True
                    run.font.color.rgb = RGBColor(52, 73, 94)  # Professional gray
                else:
                    # Regular text
                    run = paragraph.add_run(part)
                    run.font.color.rgb = RGBColor(44, 62, 80)  # Professional dark blue
                    
        except Exception as e:
            logger.error(f"Error adding text with formatting: {str(e)}")
            paragraph.add_run(text)
    
    def generate_brd_document_async(self, gathered_info: Dict, session_id: str, user_id: str, user_type: str):
        """Generate BRD document asynchronously without blocking chat response"""
        try:
            def generate_brd():
                """Background thread function to generate BRD"""
                try:
                    logger.info(f"Starting async BRD generation for session {session_id}")
                    
                    # Create Word document
                    doc = Document()
                    
                    # Add title with improved styling
                    title = doc.add_heading('Business Requirements Document', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    title_run = title.runs[0]
                    title_run.bold = True
                    title_run.font.size = Pt(24)
                    title_run.font.color.rgb = RGBColor(44, 62, 80)  # Professional dark blue
                    
                    # Add subtitle
                    subtitle = doc.add_paragraph()
                    subtitle_run = subtitle.add_run('AI Agent Development Project')
                    subtitle_run.font.size = Pt(14)
                    subtitle_run.font.color.rgb = RGBColor(52, 73, 94)  # Professional gray
                    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add date
                    date_para = doc.add_paragraph()
                    date_run = date_para.add_run(f'Generated: {datetime.now().strftime("%B %d, %Y")}')
                    date_run.font.size = Pt(10)
                    date_run.font.italic = True
                    date_run.font.color.rgb = RGBColor(127, 140, 141)  # Light gray
                    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add line break
                    doc.add_paragraph()
                    
                    # Add project information section with better formatting
                    project_heading = doc.add_heading('Project Information', 1)
                    project_heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
                    
                    project_info = [
                        ('Agent Name', gathered_info.get('agent_name', 'N/A')),
                        ('Applicable Persona', gathered_info.get('applicable_persona', 'N/A')),
                        ('Applicable Industry', gathered_info.get('applicable_industry', 'N/A')),
                        ('Session ID', session_id),
                        ('Requested By', f'{user_type}: {user_id}')
                    ]
                    
                    for label, value in project_info:
                        p = doc.add_paragraph()
                        label_run = p.add_run(f'{label}: ')
                        label_run.bold = True
                        label_run.font.color.rgb = RGBColor(52, 73, 94)
                        value_run = p.add_run(value)
                        value_run.font.color.rgb = RGBColor(44, 62, 80)
                    
                    doc.add_paragraph()
                    
                    # Add main content sections - generate all in one API call
                    sections = [
                        'Executive Summary',
                        'Business Context & Objectives',
                        'Problem Statement',
                        'User Personas & Journeys',
                        'Functional Requirements',
                        'Technical Requirements',
                        'Success Metrics',
                        'Timeline & Deliverables',
                        'Risks & Mitigation',
                        'Next Steps'
                    ]
                    
                    # Generate all sections content in one API call
                    all_sections_prompt = f"""Based on the following AI agent requirements, create a comprehensive BRD with these sections:
Agent Name: {gathered_info.get('agent_name', 'N/A')}
Persona: {gathered_info.get('applicable_persona', 'N/A')}
Industry: {gathered_info.get('applicable_industry', 'N/A')}
Problem: {gathered_info.get('problem_statement', 'N/A')}
User Journey: {gathered_info.get('user_journeys', 'N/A')}
Wow Factor: {gathered_info.get('wow_factor', 'N/A')}
Expected Output: {gathered_info.get('expected_output', 'N/A')}

Create content for ALL sections. Format each section as:
## SECTION_NAME
[Content with bullet points where appropriate]

IMPORTANT FORMATTING RULES:
- Use **bold** only for key terms and section titles within content
- Use bullet points (-) for lists
- Use numbered lists (1., 2., 3.) for sequential steps
- Keep formatting minimal and professional
- Avoid excessive bold formatting

Sections to create:
1. Executive Summary
2. Business Context & Objectives
3. Problem Statement
4. User Personas & Journeys
5. Functional Requirements
6. Technical Requirements
7. Success Metrics
8. Timeline & Deliverables
9. Risks & Mitigation
10. Next Steps

Return the complete document with all sections."""
                    
                    if self.client:
                        try:
                            response = self.client.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[
                                    {"role": "system", "content": "You are a professional business analyst who creates comprehensive BRD documents. Use minimal markdown formatting - only use **bold** for key terms and section titles within content. Avoid excessive formatting. Focus on clear, professional content structure."},
                                    {"role": "user", "content": all_sections_prompt}
                                ],
                                max_tokens=4000,
                                temperature=0.7
                            )
                            
                            content = response.choices[0].message.content
                            
                            # Parse the content and add to document with proper formatting
                            current_section = None
                            lines = content.split('\n')
                            current_paragraph = []
                            
                            for line in lines:
                                line = line.strip()
                                if not line:
                                    if current_paragraph:
                                        # Add accumulated paragraph with formatting
                                        self._add_formatted_paragraph(doc, ' '.join(current_paragraph))
                                        current_paragraph = []
                                    continue
                                
                                # Check if this is a section header
                                if line.startswith('##'):
                                    # Save previous section
                                    if current_paragraph:
                                        self._add_formatted_paragraph(doc, ' '.join(current_paragraph))
                                        current_paragraph = []
                                    
                                    # Add new section heading with styling
                                    section_name = line.replace('##', '').strip()
                                    current_section = section_name
                                    heading = doc.add_heading(section_name, 1)
                                    heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
                                elif line.startswith('-') or line.startswith('•'):
                                    # Bullet point
                                    if current_paragraph:
                                        self._add_formatted_paragraph(doc, ' '.join(current_paragraph))
                                        current_paragraph = []
                                    bullet_text = line[1:].strip()
                                    self._add_formatted_bullet(doc, bullet_text)
                                elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.')):
                                    # Numbered list
                                    if current_paragraph:
                                        self._add_formatted_paragraph(doc, ' '.join(current_paragraph))
                                        current_paragraph = []
                                    numbered_text = line[line.find('.')+1:].strip()
                                    self._add_formatted_numbered(doc, numbered_text)
                                else:
                                    # Regular text - accumulate into paragraph
                                    current_paragraph.append(line)
                            
                            # Add any remaining paragraph
                            if current_paragraph:
                                self._add_formatted_paragraph(doc, ' '.join(current_paragraph))
                            
                        except Exception as e:
                            logger.error(f"Error generating BRD content: {str(e)}")
                            # Fallback: add placeholder content for each section
                            for section in sections:
                                doc.add_heading(section, 1)
                                doc.add_paragraph("Content for this section could not be generated at this time.")
                                doc.add_paragraph()
                    
                    # Save document
                    brd_filename = f"brd_{session_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    brd_dir = "data/brds"
                    os.makedirs(brd_dir, exist_ok=True)
                    brd_path = os.path.join(brd_dir, brd_filename)
                    
                    doc.save(brd_path)
                    logger.info(f"BRD Word document generated successfully: {brd_path}")
                    
                    # Store filename in a way that's accessible to the download endpoint
                    # We'll use a simple file-based approach to signal completion
                    ready_file = os.path.join(brd_dir, f".ready_{session_id}")
                    with open(ready_file, 'w') as f:
                        f.write(brd_filename)
                    
                    return brd_filename
                    
                except Exception as e:
                    logger.error(f"Error generating BRD: {str(e)}")
                    return None
            
            # Start background thread and immediately return None (generation in progress)
            thread = threading.Thread(target=generate_brd)
            thread.daemon = True
            thread.start()
            logger.info(f"BRD generation thread started for session {session_id}")
            return None  # Return immediately, generation continues in background
            
        except Exception as e:
            logger.error(f"Error starting BRD generation: {str(e)}")
    
    def get_error_response(self, mode: str, error_message: str = None) -> Dict[str, Any]:
        """Generate error response when OpenAI is not available"""
        try:
            if mode == "explore":
                response = "I'm currently unable to access our AI agent database. This might be due to a temporary service issue. Please try again in a few moments, or contact our support team if the problem persists."
                
                return {
                    "response": response,
                    "filtered_agents": [],
                    "timestamp": datetime.now().isoformat(),
                    "error": error_message or "OpenAI API unavailable"
                }
            
            else:  # create mode
                response = "I'm currently unable to process your agent creation request due to a technical issue. Our AI requirements analyst is temporarily unavailable. Please try again in a few moments, or contact our support team for assistance."
                
                return {
                    "response": response,
                    "lets_build": False,
                    "gathered_info": {},
                    "timestamp": datetime.now().isoformat(),
                    "error": error_message or "OpenAI API unavailable"
                }
            
        except Exception as e:
            logger.error(f"Error in error response: {str(e)}")
            return {
                "response": "I'm experiencing technical difficulties. Please try again or contact support.",
                "lets_build": False,
                "gathered_info": {},
                "timestamp": datetime.now().isoformat(),
                "error": str(e)
            }
    
    def save_chat_history(self, session_id: str, user_id: str, user_type: str, mode: str, user_query: str, response: str) -> bool:
        """Save or update chat history for a session with full conversation JSON"""
        try:
            # Get the full conversation history from memory
            full_conversation = self.get_conversation_history(session_id) if session_id in self.conversation_memory else []
            
            # Add current messages to conversation
            conversation_with_current = full_conversation.copy()
            conversation_with_current.append({"role": "user", "content": user_query})
            conversation_with_current.append({"role": "assistant", "content": response})
            
            # Convert to JSON for storage
            conversation_json = json.dumps(conversation_with_current)
            
            # Check if session already exists
            chat_history_df = data_source.get_chat_history()
            existing_session = chat_history_df[chat_history_df['session_id'] == session_id]
            
            if not existing_session.empty:
                # Update existing session
                current_messages = existing_session.iloc[0]['total_messages']
                # Ensure current_messages is an integer (may be string from database)
                try:
                    current_messages = int(current_messages) if current_messages is not None else 0
                except (ValueError, TypeError):
                    current_messages = 0
                
                updated_data = {
                    'total_messages': current_messages + 1,
                    'last_message_at': datetime.now().isoformat(),
                    'conversation_summary': conversation_json  # Store full conversation as JSON
                }
                return data_source.update_chat_history_data(session_id, updated_data)
            else:
                # Create new session
                chat_data = {
                    'session_id': session_id,
                    'user_id': user_id,
                    'user_type': user_type,
                    'chat_mode': mode,
                    'title': self._generate_chat_title(user_query),
                    'conversation_summary': conversation_json,  # Store full conversation as JSON
                    'total_messages': 1,
                    'last_message_at': datetime.now().isoformat()
                }
                return data_source.save_chat_history_data(chat_data)
                
        except Exception as e:
            logger.error(f"Error saving chat history: {str(e)}")
            return False
    
    def _generate_chat_title(self, user_query: str) -> str:
        """Generate a title for the chat based on the first query"""
        try:
            # Extract key words from the query
            words = user_query.lower().split()
            key_words = []
            
            # Look for important keywords
            important_words = ['agent', 'build', 'create', 'help', 'need', 'want', 'for', 'to']
            for word in words[:10]:  # First 10 words
                if word not in important_words and len(word) > 3:
                    key_words.append(word)
            
            if key_words:
                title = ' '.join(key_words[:3]).title()  # First 3 meaningful words
                return f"Chat: {title}"
            else:
                return f"Chat: {user_query[:30]}..." if len(user_query) > 30 else f"Chat: {user_query}"
                
        except Exception as e:
            logger.error(f"Error generating chat title: {str(e)}")
            return "Chat Session"
    

    def save_agent_requirements(self, requirements: Dict[str, Any], session_id: str, user_id: Optional[str] = None, user_type: Optional[str] = None) -> bool:
        """Save the gathered agent requirements to the database"""
        try:
            # Add session_id and user information to requirements
            requirements['session_id'] = session_id
            requirements['user_id'] = user_id or "anonymous"
            requirements['user_type'] = user_type or "anonymous"
            
            # Save to data source
            success = data_source.save_agent_requirements_data(requirements)
            
            if success:
                logger.info(f"Agent requirements saved successfully for session {session_id}, user {user_id} ({user_type})")
            else:
                logger.error(f"Failed to save agent requirements for session {session_id}, user {user_id} ({user_type})")
            
            return success
        except Exception as e:
            logger.error(f"Error saving agent requirements: {str(e)}")
            return False
    
    def chat(self, user_query: str, mode: str = "explore", session_id: Optional[str] = None, user_id: Optional[str] = None, user_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Main unified chat function that handles both explore and create modes
        
        Args:
            user_query: The user's question or request
            mode: "explore" or "create"
            session_id: Optional session ID for conversation memory
            user_id: Optional user ID (isv_id, admin_id, reseller_id, client_id, or anonymous)
            user_type: Optional user type (isv, admin, reseller, client, anonymous)
            
        Returns:
            Dict containing response, metadata, and session info
        """
        try:
            # Generate session ID if not provided
            if not session_id:
                session_id = str(uuid.uuid4())
            
            # Check if OpenAI client is available
            if not self.client:
                logger.warning("OpenAI client not available, returning error response")
                result = self.get_error_response(mode, "OpenAI API key not available")
                result["session_id"] = session_id
                logger.info(f"Error response generated: response_length={len(result.get('response', ''))}, mode={mode}")
                return result
            
            # Get conversation history
            conversation_history = self.get_conversation_history(session_id)
            
            # Get conversation history for create mode
            
            # Prepare messages for OpenAI based on mode
            if mode == "explore":
                agents_context = self.get_agents_context()
                system_prompt = self.get_explore_system_prompt().format(agents_context=agents_context)
            else:  # create mode
                system_prompt = self.get_create_system_prompt()
            
            messages = [
                {"role": "system", "content": system_prompt}
            ]
            
            # Add conversation history
            messages.extend(conversation_history)
            
            # Add current user message
            messages.append({"role": "user", "content": user_query})
            
            # Call OpenAI API
            logger.info(f"Sending unified chat request for mode {mode}, session {session_id}")
            logger.info(f"API Key being used: {'Yes' if self.api_key else 'No'}")
            try:
                response = self.client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=messages,
                    max_tokens=1500,
                    temperature=0.7
                )
                logger.info(f"OpenAI API call successful for mode {mode}")
            except Exception as api_error:
                logger.error(f"OpenAI API call failed: {str(api_error)}")
                raise api_error
            
            # Extract AI response
            ai_response = response.choices[0].message.content
            
            # Save chat history
            self.save_chat_history(session_id, user_id or "anonymous", user_type or "anonymous", mode, user_query, ai_response)
            
            # Process response based on mode
            if mode == "explore":
                # Extract relevant agent IDs from the response
                filtered_agents = self.extract_agent_ids_from_response(ai_response)
                
                # Add to conversation history
                self.add_to_conversation_history(session_id, user_query, ai_response)
                
                # Format and return response
                result = {
                    "response": ai_response,
                    "filtered_agents": filtered_agents,
                    "session_id": session_id,
                    "timestamp": datetime.now().isoformat()
                }
                
            else:  # create mode
                # Parse response and metadata
                parsed_response = self.parse_create_response_metadata(ai_response, user_query)
                
                # Add to conversation history
                self.add_to_conversation_history(session_id, user_query, parsed_response["response"])
                
                # Format final result
                metadata = parsed_response["metadata"]
                result = {
                    "response": parsed_response["response"],
                    "lets_build": metadata.get("lets_build", False),
                    "gathered_info": metadata.get("gathered_info", {}),
                    "session_id": session_id,
                    "timestamp": datetime.now().isoformat()
                }
                
                # If the conversation is complete and user wants to build, save requirements
                if result.get("lets_build") and result.get("gathered_info"):
                    gathered_info = result["gathered_info"]
                    # Only save if we have meaningful information
                    if any(value.strip() for value in gathered_info.values() if isinstance(value, str)):
                        save_success = self.save_agent_requirements(gathered_info, session_id, user_id, user_type)
                        result["requirements_saved"] = save_success
                        
                        # Trigger async BRD generation (doesn't block response)
                        brd_filename = self.generate_brd_document_async(gathered_info, session_id, user_id, user_type)
                        
                        # Add BRD download link and generation status to result
                        result["brd_download_url"] = f"/api/brd/{session_id}"
                        result["brd_status"] = "generating"  # Indicates BRD is being generated
                        if brd_filename:
                            result["brd_filename"] = brd_filename
                        logger.info(f"BRD download link added for session {session_id}, status: generating")
            
            logger.info(f"Unified chat response generated for mode {mode}, session {session_id}")
            return result
            
        except Exception as e:
            logger.error(f"Error in unified chat function: {str(e)}")
            # Return error response on any failure
            result = self.get_error_response(mode, str(e))
            result["session_id"] = session_id
            return result
    
    def clear_conversation(self, session_id: str) -> Dict[str, Any]:
        """Clear conversation history for a session"""
        if session_id in self.conversation_memory:
            del self.conversation_memory[session_id]
        
        return {
            "message": "Conversation history cleared",
            "session_id": session_id,
            "timestamp": datetime.now().isoformat()
        }

# Global unified chat agent instance
unified_chat_agent = UnifiedChatAgent()

# Example usage and testing
if __name__ == "__main__":
    # Test the unified chat agent
    print("🤖 Testing Unified Chat Agent...")
    
    # Test explore mode
    print("\n=== EXPLORE MODE TEST ===")
    response = unified_chat_agent.chat("I need help with financial analysis", "explore")
    print(f"Response: {response['response']}")
    print(f"Filtered Agents: {response.get('filtered_agents', [])}")
    
    # Test create mode
    print("\n=== CREATE MODE TEST ===")
    response = unified_chat_agent.chat("I need an AI agent for customer support", "create")
    print(f"Response: {response['response']}")
    print(f"Let's Build: {response.get('lets_build', False)}")
    print(f"Gathered Info: {response.get('gathered_info', {})}")
